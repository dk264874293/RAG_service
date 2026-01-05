'''
Author: 汪培良 rick_wang@yunquna.com
Date: 2026-01-04 18:23:35
LastEditors: 汪培良 rick_wang@yunquna.com
LastEditTime: 2026-01-05 20:47:16
FilePath: /RAG_service/loader/work_loader.py
Description: 这是默认设置,请设置`customMade`, 打开koroFileHeader查看配置 进行设置: https://github.com/OBKoro1/koro1FileHeader/wiki/%E9%85%8D%E7%BD%AE
'''
#导入 Python 内置日志模块，用于记录程序运行中的日志（错误、警告、信息等）
import logging
#Python 内置 MIME 类型模块，用于通过文件内容类型 / 后缀猜测文件格式（如图片的 jpg/png 格式）
import mimetypes
#导入 Python 内置操作系统交互模块，用于处理文件路径、目录操作等
import os
#导入 Python 内置正则表达式模块，用于匹配和提取字符串（此处用于匹配超链接 URL）
import re
#导入 Python 内置临时文件模块，用于创建临时文件存储网络下载的 docx 文件，避免持久化占用磁盘
import tempfile
#导入 Python 内置唯一 ID 生成模块，用于生成唯一的图片文件名，避免文件重名覆盖
import uuid
#从 Python 内置 URL 解析模块中导入urlparse方法，用于校验 URL 是否有效
from urllib.parse import urlparse
#导入 Python 内置 XML 解析模块，用于解析 docx 中的超链接 XML 结构，提取 URL 地址
from xml.etree import ElementTree

#导入 Python 内置 XML 解析模块，用于解析 docx 中的超链接 XML 结构，提取 URL 地址
import httpx
# 导入第三方 HTTP 请求库，用于发送网络请求（下载网络 docx 文件、外部图片）
from docx import Document as DocxDocument

# 内置方法
# from configs import dify_config
# from core.helper import ssrf_proxy
from loader.extractor_base import BaseExtractor
from loader.models import Document
# from extensions.ext_database import db
# from extensions.ext_storage import storage
# from libs.datetime_utils import naive_utc_now
# from models.enums import CreatorUserRole
# from models.model import UploadFile

logger = logging.getLogger(__name__)

class WorkExtractor(BaseExtractor):
    """
    加载docx文件。
    Args:
        file_path (str): 要加载的文件路径
    """
    def __init__(self, file_path:str, tenant_id:str,user_id:str):
        """通过文件路径初始化实例"""
        self.file_path = file_path
        self.tenant_id = tenant_id
        self.user_id = user_id
        
        # 条件判断，检查实例属性self.file_path中是否包含~（Python 中表示用户主目录的简写）
        if "~" in self.file_path:
            # 若路径包含~，调用os.path.expanduser方法将~展开为绝对用户目录路径（如~展开为/Users/wangpeiliang）
            self.file_path = os.path.expanduser(self.file_path)
            
        # 条件判断，检查实例属性self.file_path是否为本地文件路径（通过os.path.isfile判断）,_is_valid_url 判断是否为有效的url
        if not os.path.isfile(self.file_path) and self._is_valid_url(self.file_path):
            response = httpx.get(self.file_path, timeout=None)
            if response.status_code != 200:
                response.close()
                raise ValueError(f"检查你的文件URL，文件下载失败，状态码：{response.status_code}")
            self.web_path = self.file_path
            # 创建临时文件，默认关闭后自动删除
            self.temp_file = tempfile.NamedTemporaryFile()
           
            try:
                self.temp_file.write(response.content)
            finally:
                response.close()
            self.file_path = self.temp_file.name
            print(f'response --> {self.file_path}')
        elif not os.path.isfile(self.file_path):
            raise ValueError(f"检查你的文件路径，{self.file_path}不是一个有效的url")
        

    def __del__(self):
        pass

    def extract(self) -> list[Document]:
        content = self.parse_docx(self.file_path)
        print(f'content --> {content}')
        return [
            Document(
                page_content=content,
                metadata={
                    "source":self.file_path,
                }
            )
        ]

    @staticmethod
    def _is_valid_url(url:str) -> bool:
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    
    def _extract_images_from_docx(self, doc):
        """图片提取方法"""
        image_count = 0
        image_map = {}
        #获取文件预览基础 URL
        base_url = "test"
        for r_id,rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                image_count += 1
                if rel.is_external:
                    url = rel.target_ref
                    if not self._is_valid_url(url):
                        continue
                    try:
                        response = requests.get(url)
                        # response.raise_for_status()
                    except Exception as e:
                        logger.warning("Failed to download image from URL: %s: %s", url, str(e))
                        continue
                    if response.status_code == 200:
                        image_ext = mimetypes.guess_extension(response.headers.get("Content-Type",""))
                        if image_ext is None:
                            continue
                        file_uuid = str(uuid.uuid4())
                        file_key = "image_files/" + self.tenant_id + "/" + file_uuid + image_ext
                        mime_type, _ = mimetypes.guess_type(file_key)
                        # TODO 待增加保存逻辑
                        print(f"response --> {response.content}")
                        image_map[r_id] = f"![image]({base_url}/files/{file_key}/file-preview)"
        return image_map
                        # storage.save(file_key,response.content)
        

    def _table_to_markdown(self,table,image_map):
        """表格转 Markdown"""
        markdown = []

        total_cols = max(len(row.cells) for row in table.rows)
        header_row = table.rows[0]
        headers = self._parse_row(header_row, image_map, total_cols)
        markdown.append("| " + " | ".join(headers) + " |")
        markdown.append("| " + " | ".join(["---"] * total_cols) + " |")

        for row in table.rows[1:]:
            row_cells = self._parse_row(row, image_map, total_cols)
            markdown.append("| " + " | ".join(row_cells) + " |")
        return "\n".join(markdown)

    def _parse_row(self,row,image_map,total_cols):
        """解析表格行"""
        row_cells= [""] * total_cols
        col_index = 0
        while col_index < len(row.cells):
            while col_index < len(row.cells) and row_cells[col_index] != "":
                col_index += 1
            if col_index >= len(row.cells):
                break
            cell = row.cells[col_index]
            cell_content = self._parse_cell(cell, image_map).strip()
            cell_colspan = cell.grid_span or 1
            for i in range(cell_colspan):
                if col_index + i < total_cols:
                    row_cells[col_index + i] = cell_content if i == 0 else ""
            col_index += cell_colspan
        return row_cells

    def _parse_cell(self,cell,image_map):
        """解析表格单元格"""
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_cell_paragraph(paragraph, image_map)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)

    def _parse_cell_paragraph(self, paragraph, image_map):
        """解析单元格段落"""
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    image_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if not image_id:
                        continue
                    rel = paragraph.part.rels.get(image_id)
                    if rel is None:
                        continue
                    if rel.is_external:
                        if image_id in image_map:
                            paragraph_content.append(image_map[image_id])
                    else:
                        image_part = rel.target_part
                        if image_part in image_map:
                            paragraph_content.append(image_map[image_part])
            else:
                paragraph_content.append(run.text)
        return "".join(paragraph_content).strip()

    def parse_docx(self, docx_path):
        """解析 docx 文件"""
        doc = DocxDocument(docx_path)

        content = []

        image_map = self._extract_images_from_docx(doc)

        hyperlinks_url = None
        url_pattern = re.compile(r"http://[^\s+]+//|https://[^\s+]+")
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and hyperlinks_url:
                     # 若有文本和临时超链接，拼接为Markdown超链接格式
                    result = f"  [{run.text}]({hyperlinks_url})  "
                    run.text = result
                    hyperlinks_url = None
                if "HYPERLINK" in run.element.xml:
                    try:
                        xml = ElementTree.XML(run.element.xml)
                        x_child = [c for c in xml.iter() if c is not None]
                        for x in x_child:
                            if x.tag.endswith("instrText") and x.text:
                                for i in url_pattern.findall(x.text):
                                    hyperlinks_url = str(i)
                    except Exception:
                        logger.exception("Failed to parse HYPERLINK xml")
        def parse_paragraph(paragraph):
            """解析段落"""
            paragraph_content = []

            def append_image_link(image_id, has_drawing):
                pass

            for run in paragraph.runs:
                if hasattr(run.element,"tag") and isinstance(run.element.tag,str) and run.element.tag.endswith('r'):
                    # Word 2007 及以上版本的图片
                    drawing_elements = run.element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    has_drawing = False
                    for drawing in drawing_elements:
                        blip_elements = drawing.findall(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        )
                        for blip in blip_elements:
                            embed_id = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed_id:
                                rel = doc.part.rels.get(embed_id)
                                if rel is not None and rel.is_external:
                                    # External image: use embed_id as key
                                    if embed_id in image_map:
                                        has_drawing = True
                                        paragraph_content.append(image_map[embed_id])
                                else:
                                    # Internal image: use target_part as key
                                    image_part = doc.part.related_parts.get(embed_id)
                                    if image_part in image_map:
                                        has_drawing = True
                                        paragraph_content.append(image_map[image_part])
                    
                    shape_elements = run.element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
                    )
                    for shape in shape_elements:
                        shape_image = shape.find(
                            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}binData"
                        )
                        if shape_image is not None and shape_image.text:
                            image_id = shape_image.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                            )
                            if image_id and image_id in doc.part.rels:
                                append_image_link(image_id, has_drawing)
                        image_data = shape.find(".//{urn:schemas-microsoft-com:vml}imagedata")
                        if image_data is not None:
                            image_id = image_data.get("id") or image_data.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                            )
                            if image_id and image_id in doc.part.rels:
                                append_image_link(image_id, has_drawing)
                if run.text.strip():
                    paragraph_content.append(run.text.strip())
            return "".join(paragraph_content) if paragraph_content else ""
            pass
        
        paragraphs = doc.paragraphs.copy()
        tables = doc.tables.copy()
        for element in doc.element.body:
            if hasattr(element,"tag"):
                if isinstance(element.tag, str) and element.tag.endswith("p"):
                    para = paragraphs.pop(0)
                    parsed_paragraph = parse_paragraph(para)
                    if parsed_paragraph.strip():
                        content.append(parsed_paragraph)
                    else:
                        content.append("\n")
                elif isinstance(element.tag,str) and element.tag.endswith('tbl'):
                    table = tables.pop(0)
                    content.append(self._table_to_markdown(table,image_map))

        return '\n'.join(content)

        
